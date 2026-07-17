"""Word (.docx) loader: paragraphs plus tables, flattened to plain text."""

from __future__ import annotations

from pathlib import Path

from graphrag.core.errors import IngestionError
from graphrag.core.types import Document
from graphrag.ingestion.loaders.base import Loader


class DocxLoader(Loader):
    suffixes = (".docx",)

    def load(self, path: Path) -> Document:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise IngestionError("python-docx is not installed") from exc

        doc = docx.Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            ]
            if rows:
                parts.append("\n".join(rows))
        return Document(
            source=str(path), content="\n\n".join(parts), metadata={"type": "docx"}
        )
